from flask import Flask, request, jsonify, redirect, render_template
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import requests
import json
import os
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

@app.route('/')
def home():
    #return render_template('index.html')
    return jsonify({"message": "Welcome to the Flask API!"}), 200

# 이미지 저장 경로 설정
UPLOAD_FOLDER = 'uploads'  # 업로드할 폴더 이름
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)  # 폴더가 없으면 생성

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/upload-image', methods=['POST'])
def upload_image():
    # 사용자가 업로드한 파일 가져오기
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']

    # 파일이 없거나 파일 이름이 비어있을 경우
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    # 파일 저장
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    return jsonify({"message": "Image uploaded successfully", "file_path": file_path}), 200

CHAT_API_URL = "https://flow.edutrack.kr/api/v1/prediction/35de7e76-2bdc-4fb9-87ac-24c618dcbcad"
SUM_API_URL = "https://flow.edutrack.kr/api/v1/prediction/820a4358-23bd-4abb-aa1e-4c7dad121e90"
JOB_API_URL = "https://flow.edutrack.kr/api/v1/prediction/0beb8cf8-af68-4779-bfdd-bde54737215a"

def chat_query(payload):
    response = requests.post(CHAT_API_URL, json=payload, verify=False)
    return response.json()

def sum_query(payload):
    response = requests.post(SUM_API_URL, json=payload, verify=False)
    return response.json()

chat_output = {}
last_chat_output = {}

@app.route('/chat', methods=['POST'])
def chat():
    global last_chat_output
    global chat_output
    print('Activate chat')
    print('Received request:', request.json)

    user_input = request.json.get('question')
    
    if not user_input:
        return jsonify({"error": "No question provided"}), 400
    
    if user_input == 'exit':
        last_chat_output = chat_output

    chat_output = chat_query({'question': user_input})
    print('chat', chat_output)
    return jsonify(chat_output)

def summarize(last_chat):
    if not last_chat:
        return {"error": "Invalid input format"}, 400

    sum_output = sum_query({"question": last_chat})
    sum_txt = sum_output['text'].replace('```', '').replace('json\n', '')

    try:
        json_data = json.loads(sum_txt)
    except json.JSONDecodeError as e:
        return {"error": "Failed to decode JSON", "details": str(e)}

    return json_data

def load_image():
    image_path = 'picture.jpg'
    return image_path

def add_image(table, image_path):
    img_1 = table.rows[0].cells[0]
    img_2 = table.rows[1].cells[0]
    img_3 = table.rows[2].cells[0]
    img_4 = table.rows[3].cells[0]

    merged_cell = img_1.merge(img_2).merge(img_3).merge(img_4)

    for paragraph in merged_cell.paragraphs:
        paragraph.clear()

    run = merged_cell.add_paragraph().add_run()
    run.add_picture(image_path, width=Inches(1))

def add_txt(table, row, col, data):
    if data == 0 or data == []:
        data = '-'
    table.rows[row].cells[col].text = data
    table.rows[row].cells[col].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def documentation(json_data):

    # json에서 유저 정보 불러오기
    kor_name = json_data['kor_name']
    eng_name = json_data['eng_name']
    phone_no = json_data['phone_no']
    email = json_data['email']
    cnt_address = json_data['cnt_address']

    # 이력서 양식 불러오기
    doc_path = 'templates/work_01.docx'
    doc = Document(doc_path)
    table = doc.tables[0]

    # 기본정보
    add_image(table, 'picture.jpg')
    add_txt(table, 1,3,kor_name)
    add_txt(table, 2,3,eng_name)
    add_txt(table, 2,5,phone_no)
    add_txt(table, 3,3,email)
    add_txt(table, 4,1,cnt_address)

    education = json_data['education']
    for i, e in enumerate(education):
        edu_duration = e['duration']
        edu_school = e['school_name']
        edu_major = e['major']
        edu_note = e['note']

        tmp_r = 7+i
        add_txt(table, tmp_r,1,edu_duration)
        add_txt(table, tmp_r,3,edu_school)
        add_txt(table, tmp_r,5,edu_major)
        add_txt(table, tmp_r,6,edu_note)

    prac_exp = json_data['prac_exp']
    for i, e in enumerate(prac_exp):
        prac_duration = e['duration']
        prac_place = e['workplace']
        prac_duty = e['duty']
        prac_retire = e['retire_reason']

        tmp_r = 12+i
        add_txt(table, tmp_r,1,prac_duration)
        add_txt(table, tmp_r,3,prac_place)
        add_txt(table, tmp_r,5,prac_duty)
        add_txt(table, tmp_r,6,prac_retire)

    other_exp = json_data['other_exp']
    for i, e in enumerate(other_exp):
        other_duration = e['duration']
        other_place = e['workplace']
        other_duty = e['duty']
        other_retire = e['retire_reason']

        tmp_r = 19+i
        add_txt(table, tmp_r,1,other_duration)
        add_txt(table, tmp_r,3,other_place)
        add_txt(table, tmp_r,5,other_duty)
        add_txt(table, tmp_r,6,other_retire)

    now = datetime.now()
    add_txt(table, 24,0,f'위 기재 사항은 사실과 틀림이 없습니다.\n\n{now.strftime("%Y년   %m월   %d일")}\n\n지원자 :   {kor_name}   (인)')

    new_doc_path = 'new_work_01.docx'
    doc.save(new_doc_path)

@app.route('/documentation', methods=['POST'])
def generate_document():
    try:
        user_data = '''
            알겠습니다! 작성된 이력서를 정리해서 보여드리겠습니다.

            이력서 (RESUME)
            한글 이름 (Kor Name): 응우옌 타이 마이
            영어 이름 (Eng Name): Nguyen Thai Mai
            전화번호 (Phone No): 010-1234-1234
            이메일 (Email):
            exam1234@gmail.com
            주소 (Cnt Address): 한국시 한국도 한국동 한국아파트 111-102호
            학력 (Education)
            기간 (Duration): 2012년 3월 - 2016년 3월
            학교 이름 (School Name): Trường Trung Học Phổ Thông Nguyễn Thị Minh Khai
            근무 경력 (Practical Experience)
            기간 (Duration): 2017년
            회사 이름 (Workplace): Shop Quần Áo Đẹp
            직무 (Duty): 옷 판매
            퇴사 이유 (Retire Reason): 한국 이주 후 일하지 않음
            추가 정보 (Additional Information)
            베트남에서 한국어 학습: 베트남에서 한국어를 공부함
            이렇게 이력서 작성을 완료했습니다! 다른 도움이 필요하시면 언제든지 말씀해 주세요. 이 세션을 종료하시려면 "exit"을 입력해 주세요.
            '''

        json_data = summarize(user_data)
        print('json', json_data)

        documentation(json_data)
        return jsonify({"message": "Document created successfully", "file_path": "new_work_01.docx"}), 200
    except Exception as e:
        return jsonify({"error": "Failed to create document", "details": str(e)}), 500
    
def job_query(payload):
    response = requests.post(JOB_API_URL, json=payload, verify=False)
    return response.json()

@app.route('/job', methods=['POST'])
def job_chat():
    user_input = request.json.get('question')
    job_output = job_query({
        'question' : user_input
    })
    print(job_output)
    return jsonify(job_output)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)